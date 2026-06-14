import json
import os

transcript_path = r'C:\Users\muhdz\.gemini\antigravity\brain\14706109-6b48-4677-8a00-de64a460451e\.system_generated\logs\transcript_full.jsonl'
docx_path = r'C:\Users\muhdz\Downloads\Documents\Project\PRESENSE_Product_Specification.docx'

out_path = r'C:\Users\muhdz\.gemini\antigravity\scratch\presense\output.txt'

with open(out_path, 'w', encoding='utf-8') as out_f:
    out_f.write("--- DIAGNOSTIC FROM LOGS ---\n")
    with open(transcript_path, 'r', encoding='utf-8') as f:
        for line in f:
            if '"type":"USER_INPUT"' in line and 'Full Diagnostic' in line:
                data = json.loads(line)
                content = data.get('content', '')
                if 'Full Diagnostic' in content:
                    out_f.write(content)
                    break

    out_f.write("\n--- DOCX TEXT ---\n")
    try:
        import docx
        doc = docx.Document(docx_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        out_f.write('\n'.join(fullText))
    except ImportError:
        import subprocess
        subprocess.run(['pip', 'install', 'python-docx'])
        import docx
        doc = docx.Document(docx_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        out_f.write('\n'.join(fullText))
    except Exception as e:
        out_f.write(f"Could not read docx: {e}\n")
